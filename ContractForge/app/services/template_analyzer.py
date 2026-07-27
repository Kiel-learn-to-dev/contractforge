"""
template_analyzer.py — Phân tích file .docx, trích xuất placeholder và
xây dựng field mapping mặc định.

Hai thách thức kỹ thuật chính của python-docx:
  1. SPLIT RUN: Word lưu text trong nhiều run riêng trong cùng một paragraph
     (do autocorrect, formatting, tracked changes). Ví dụ {{CONTRACT_NUMBER}}
     có thể bị lưu thành 3 run: '{{', 'CONTRACT', '_NUMBER}}'.
     → Giải pháp: join toàn bộ run.text trong para TRƯỚC khi regex.

  2. VỊ TRÍ: placeholder có thể nằm ở para thường, trong cell bảng,
     header/footer của từng section.
     → Phải quét đủ 3 chỗ.

Hàm công khai:
  extract_placeholders(path_or_bytes) → list[str]   (sorted, deduped)
  build_default_mapping(placeholders) → dict[str,str]
  analyze(path_or_bytes) → AnalysisResult
"""

import io
import re
from dataclasses import dataclass, field
from typing import Union
import docx


# ─── Canonical field registry (khớp với blueprint §field canonical) ──────────

CANONICAL_MAP: dict[str, str] = {
    # Hợp đồng
    "CONTRACT_NUMBER":          "contract.number",
    "SIGN_DATE_DAY":            "contract.sign_date_day",
    "SIGN_DATE_MONTH":          "contract.sign_date_month",
    "SIGN_DATE_YEAR":           "contract.sign_date_year",
    "START_DATE":               "contract.start_date",
    "END_DATE":                 "contract.end_date",
    "MONTH_COUNT":              "contract.month_count",

    # Giá trị
    "UNIT_COUNT":               "contract.unit_count",
    "QUANTITY":                 "contract.quantity",
    "UNIT_PRICE_PRE_VAT":       "contract.unit_price_pre_vat",
    "VAT_RATE":                 "contract.vat_rate",
    "UNIT_PRICE_VAT":           "contract.unit_price_vat",
    "TOTAL_AMOUNT":             "contract.total_amount",
    "AMOUNT_TEXT":              "contract.amount_text",
    "AMOUNT_TEXT_CHAN":         "contract.amount_text_chan",
    "PRICING_DESCRIPTION":      "contract.pricing_description",
    "PRICING_UNITS":            "quotation.pricing_units",

    # Báo giá
    "QUOTE_DATE":               "quotation.date",
    "QUOTE_DATE_DAY":           "quotation.date_day",
    "QUOTE_DATE_MONTH":         "quotation.date_month",
    "QUOTE_DATE_YEAR":          "quotation.date_year",

    # Sản phẩm
    "PRODUCT_CODE":             "product.code",
    "PRODUCT_NAME":             "product.name",

    # Nghiệm thu
    "ACCEPTANCE_DATE":          "contract.acceptance_date",
    "ACCEPTANCE_DATE_DAY":      "contract.acceptance_date_day",
    "ACCEPTANCE_DATE_MONTH":    "contract.acceptance_date_month",
    "ACCEPTANCE_DATE_YEAR":     "contract.acceptance_date_year",

    # Thanh lý
    "LIQUIDATION_DATE_DAY":     "contract.liquidation_date_day",
    "LIQUIDATION_DATE_MONTH":   "contract.liquidation_date_month",
    "LIQUIDATION_DATE_YEAR":    "contract.liquidation_date_year",

    # Bên A (khách hàng)
    "PARTY_A_NAME":             "party_a.name",
    "PARTY_A_ADDRESS":          "party_a.address",
    "PARTY_A_REPRESENTATIVE":   "party_a.representative",
    "PARTY_A_TITLE":            "party_a.title",
    "PARTY_A_NAME_WARD":        "party_a.name_ward",   # HSSK specific
    "PARTY_A_TAX_CODE":         "party_a.tax_code",
    "PARTY_A_BANK_ACCOUNT":     "party_a.bank_account",
    "PARTY_A_BANK_NAME":        "party_a.bank_name",
    "PARTY_A_BENEFICIARY":      "party_a.beneficiary",
    "PARTY_A_BUDGET_CODE":      "party_a.budget_code",

    # Bên B (Viettel – cố định, nhưng đại diện có thể thay đổi)
    "PARTY_B_REPRESENTATIVE":       "party_b.representative",
    "PARTY_B_AUTHORIZATION_NUMBER": "party_b.authorization_number",
    "PARTY_B_AUTHORIZATION_DATE":   "party_b.authorization_date",
    "PARTY_B_NAME":                  "party_b.name",
    "PARTY_B_ADDRESS":               "party_b.address",
    "PARTY_B_BANK_ACCOUNT":          "party_b.bank_account",
    "PARTY_B_BANK_NAME":             "party_b.bank_name",
    "PARTY_B_BENEFICIARY":           "party_b.beneficiary",
    "PARTY_B_TAX_CODE":              "party_b.tax_code",
    "PARTY_B_TITLE":                 "party_b.title",
}

# Toàn bộ canonical fields cho UI dropdown
ALL_CANONICAL_FIELDS: list[str] = sorted(set(CANONICAL_MAP.values()))

# Regex pattern — chỉ match chữ HOA, số và dấu gạch dưới
_PLACEHOLDER_RE = re.compile(r'\{\{([A-Z][A-Z0-9_]*)\}\}')


# ─── Data types ──────────────────────────────────────────────────────────────

@dataclass
class AnalysisResult:
    placeholders: list[str] = field(default_factory=list)
    """Danh sách tên placeholder, đã sắp xếp, không trùng."""

    mapping: dict[str, str] = field(default_factory=dict)
    """placeholder name → canonical field (theo CANONICAL_MAP, có thể thiếu)."""

    split_run_warnings: list[str] = field(default_factory=list)
    """Danh sách paragraph text chứa placeholder bị split-run."""

    source_counts: dict[str, int] = field(default_factory=dict)
    """Số placeholder tìm thấy theo nguồn: para / table / header / footer."""

    error: str = ""
    """Nếu không trống → quá trình phân tích thất bại."""

    @property
    def ok(self) -> bool:
        return not self.error


# ─── Public API ──────────────────────────────────────────────────────────────

def analyze(source: Union[str, bytes]) -> AnalysisResult:
    """
    Phân tích file .docx (đường dẫn hoặc bytes).
    Trả về AnalysisResult đầy đủ.
    """
    result = AnalysisResult()
    try:
        doc = _open_doc(source)
    except Exception as e:
        result.error = f"Không mở được file .docx: {e}"
        return result

    seen: set[str] = set()
    counts = {"para": 0, "table": 0, "header": 0, "footer": 0}

    # 1. Body paragraphs
    for para in doc.paragraphs:
        ph_list = _extract_from_para(para, result.split_run_warnings)
        for ph in ph_list:
            if ph not in seen:
                seen.add(ph)
            counts["para"] += 1

    # 2. Tables (tất cả cells, tất cả paragraphs trong cell)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    ph_list = _extract_from_para(para, result.split_run_warnings)
                    for ph in ph_list:
                        if ph not in seen:
                            seen.add(ph)
                        counts["table"] += 1

    # 3. Headers & Footers (mỗi section)
    for section in doc.sections:
        for hf_name, counter_key in [("header", "header"), ("footer", "footer")]:
            hf = getattr(section, hf_name, None)
            if hf is None:
                continue
            for para in hf.paragraphs:
                ph_list = _extract_from_para(para, result.split_run_warnings)
                for ph in ph_list:
                    if ph not in seen:
                        seen.add(ph)
                    counts[counter_key] += 1

    result.placeholders = sorted(seen)
    result.mapping = build_default_mapping(result.placeholders)
    result.source_counts = counts
    return result


def extract_placeholders(source: Union[str, bytes]) -> list[str]:
    """Convenience wrapper — chỉ trả danh sách placeholder."""
    return analyze(source).placeholders


def build_default_mapping(placeholders: list[str]) -> dict[str, str]:
    """
    Tạo mapping mặc định từ CANONICAL_MAP.
    Những placeholder không có trong registry vẫn được giữ lại với
    canonical = "" (người dùng cần tự điền).
    """
    return {ph: CANONICAL_MAP.get(ph, "") for ph in placeholders}


# ─── Internals ───────────────────────────────────────────────────────────────

def _open_doc(source: Union[str, bytes]) -> docx.Document:
    if isinstance(source, bytes):
        return docx.Document(io.BytesIO(source))
    return docx.Document(source)


def _extract_from_para(para, split_warnings: list[str]) -> list[str]:
    """
    Trích xuất tất cả placeholder từ một paragraph.

    Chiến lược:
      - Trước tiên thử tìm trực tiếp từ para.text (join của tất cả runs).
      - Nếu tìm thấy, kiểm tra xem có run nào chứa '{{' nhưng không chứa '}}'
        (split run) để ghi warning.
      - Không tìm trên từng run riêng lẻ — luôn dùng full text.
    """
    full_text = para.text  # python-docx tự join runs
    found = _PLACEHOLDER_RE.findall(full_text)
    if not found:
        return []

    # Phát hiện split run để warn
    has_open = any('{{' in r.text for r in para.runs)
    if has_open:
        for run in para.runs:
            if '{{' in run.text and '}}' not in run.text:
                warning = full_text[:120].strip()
                if warning not in split_warnings:
                    split_warnings.append(warning)
                break

    return found
