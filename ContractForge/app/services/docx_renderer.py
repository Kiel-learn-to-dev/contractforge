"""
docx_renderer.py — Thay thế placeholder trong file .docx một cách an toàn.

Vấn đề cốt lõi của python-docx:
  Word lưu text thành nhiều "run" trong cùng một paragraph. Khi bạn
  gõ {{CONTRACT_NUMBER}} trong Word, đôi khi nó bị lưu thành 3-4 run
  riêng lẻ (do autocorrect, spell check, format change giữa chừng).
  Nếu chỉ loop qua từng run.text và replace, sẽ bỏ sót những trường hợp bị split.

Giải pháp (2 lớp):
  Lớp 1 — Single-run replace (nhanh, giữ nguyên format):
    Nếu {{KEY}} nằm gọn trong 1 run → replace trực tiếp run.text.

  Lớp 2 — Merge-and-replace (fallback khi split-run):
    Nếu sau lớp 1 vẫn còn {{...}} trong para.text:
      - Merge toàn bộ runs vào run đầu tiên (giữ format của run đó)
      - Replace toàn bộ placeholder trên text đã merge
      - Xóa các run thừa

Hàm công khai:
  render_docx(template_path_or_bytes, context) → bytes
  build_render_context(contract, customer, product, template) → dict[str, str]
"""

import copy
import io
import os
import re
from decimal import Decimal
from typing import Union

import docx
from docx.oxml.ns import qn

from app.utils.amount_to_words_vi import amount_to_words, amount_to_words_chan
from app.utils.date_helpers import format_date_vn, split_date_parts, parse_date_input

_PH_RE = re.compile(r'\{\{([A-Z][A-Z0-9_]*)\}\}')


# ─── Dynamic Party A row insertion ──────────────────────────────────────────

def _copy_row_format(src_row, table):
    """Copy một row từ bảng, giữ nguyên định dạng XML, trả về row element mới."""
    from lxml import etree
    import copy as _copy
    new_tr = _copy.deepcopy(src_row._tr)
    # Xóa nội dung các cell nhưng giữ properties
    for tc in new_tr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
        for p in tc.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            for r in p.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                p.remove(r)
    return new_tr


def _set_cell_text(tc_elem, text: str, copy_rpr_from=None):
    """Đặt text vào cell element, kế thừa font/size từ cell mẫu."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = tc_elem.find(f'{{{W}}}p')
    if p is None:
        from lxml import etree
        p = etree.SubElement(tc_elem, f'{{{W}}}p')
    # Xóa run cũ
    for r in p.findall(f'{{{W}}}r'):
        p.remove(r)
    from lxml import etree
    r_new = etree.SubElement(p, f'{{{W}}}r')
    if copy_rpr_from is not None:
        rpr = copy_rpr_from.find(f'{{{W}}}rPr')
        if rpr is not None:
            import copy as _copy
            r_new.insert(0, _copy.deepcopy(rpr))
    t = etree.SubElement(r_new, f'{{{W}}}t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _insert_party_a_optional_rows(doc, context: dict) -> None:
    """
    Với mỗi bảng Bên A (nhận biết qua ô đầu chứa PARTY_A_NAME sau render),
    chèn thêm dòng tùy chọn sau dòng "Chức vụ" nếu context có giá trị.

    Thứ tự dòng thêm (chỉ khi có dữ liệu):
      Mã số thuế        → PARTY_A_TAX_CODE
      Số tài khoản      → PARTY_A_BANK_ACCOUNT
      Ngân hàng         → PARTY_A_BANK_NAME
      Tên thụ hưởng     → PARTY_A_BENEFICIARY
      Mã QHNS           → PARTY_A_BUDGET_CODE
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    import copy as _copy

    # Danh sách trường tùy chọn: (label, context_key)
    OPTIONAL = [
        ('Mã số thuế',    'PARTY_A_TAX_CODE'),
        ('Số tài khoản',  'PARTY_A_BANK_ACCOUNT'),
        ('Ngân hàng',     'PARTY_A_BANK_NAME'),
        ('Tên thụ hưởng', 'PARTY_A_BENEFICIARY'),
        ('Mã QHNS',       'PARTY_A_BUDGET_CODE'),
    ]

    # Lọc các trường thực sự có giá trị
    rows_to_add = [(label, context.get(key, ''))
                   for label, key in OPTIONAL
                   if context.get(key, '').strip()]

    if not rows_to_add:
        return

    party_a_name = context.get('PARTY_A_NAME', '')

    for table in doc.tables:
        # Nhận biết bảng Bên A: ô đầu tiên của row 0 chứa placeholder {{PARTY_A_NAME}}
        # (hàm được gọi TRƯỚC khi replace placeholder nên dùng raw placeholder)
        if not table.rows:
            continue
        first_cell_text = table.rows[0].cells[0].text.strip()
        if "PARTY_A_NAME" not in first_cell_text:
            continue
        # Tìm dòng "Chức vụ" (row cuối trong khối Bên A — thường là row[-1])
        # Dùng row có text cột 0 = "Chức vụ" hoặc row cuối
        title_row_idx = len(table.rows) - 1
        for i, row in enumerate(table.rows):
            if row.cells[0].text.strip() in ('Chức vụ', 'Chức danh'):
                title_row_idx = i
                break

        title_row = table.rows[title_row_idx]
        # Template cho row mới: copy từ dòng "Đại diện" (thường là title_row - 1)
        template_row = table.rows[max(0, title_row_idx - 1)]

        tbl_elem = table._tbl
        after_tr = title_row._tr

        for label, value in rows_to_add:
            new_tr = _copy.deepcopy(template_row._tr)
            tcs = new_tr.findall(f'.//{{{W}}}tc')
            # Lấy run đầu tiên trong dòng mới để kế thừa font/size (fix Bug: cỡ chữ 14)
            _ref_run = new_tr.find(f'.//{{{W}}}r')
            if len(tcs) >= 3:
                # col 0: label, col 1: ":", col 2: value
                _set_cell_text(tcs[0], label, copy_rpr_from=_ref_run)
                _set_cell_text(tcs[1], ':', copy_rpr_from=_ref_run)
                _set_cell_text(tcs[2], value, copy_rpr_from=_ref_run)
            elif len(tcs) == 2:
                _set_cell_text(tcs[0], label, copy_rpr_from=_ref_run)
                _set_cell_text(tcs[1], value, copy_rpr_from=_ref_run)
            # Chèn sau after_tr
            after_tr.addnext(new_tr)
            # after_tr = new_tr để dòng tiếp theo chèn sau
            after_tr = new_tr


# ─── Public API ───────────────────────────────────────────────────────────────

def apply_template_mapping(context: dict[str, str], template=None) -> dict[str, str]:
    """Add custom placeholder keys from a template's canonical field mapping."""
    mapped = dict(context)
    if template is None:
        return mapped

    from app.services.template_analyzer import CANONICAL_MAP

    values_by_canonical = {
        canonical: context.get(placeholder, "")
        for placeholder, canonical in CANONICAL_MAP.items()
        if placeholder in context
    }
    for placeholder, canonical in template.field_mapping.items():
        if canonical in values_by_canonical:
            mapped[placeholder] = values_by_canonical[canonical]
    return mapped

def render_docx(
    template_source: Union[str, bytes],
    context: dict[str, str],
) -> bytes:
    """
    Thay thế tất cả {{PLACEHOLDER}} trong file .docx bằng giá trị trong context.

    Args:
        template_source: Đường dẫn file hoặc bytes nội dung .docx.
        context: Dict mapping placeholder_name → giá trị (đã là str).
                 VD: {"CONTRACT_NUMBER": "HĐ-001/2025", "PARTY_A_NAME": "..."}

    Returns:
        Bytes của file .docx đã render.

    Raises:
        ValueError nếu không mở được file.
        (Không raise khi có placeholder không có trong context — để nguyên {{KEY}})
    """
    try:
        if isinstance(template_source, bytes):
            doc = docx.Document(io.BytesIO(template_source))
        else:
            doc = docx.Document(template_source)
    except Exception as e:
        raise ValueError(f"Không mở được template: {e}")

    # Chuẩn hóa context: tất cả value là str, None → ""
    ctx = {k: ("" if v is None else str(v)) for k, v in context.items()}

    # Chèn dòng tùy chọn vào bảng Bên A trước khi thay placeholder
    _insert_party_a_optional_rows(doc, ctx)

    # Xử lý body paragraphs
    for para in doc.paragraphs:
        _render_paragraph(para, ctx)

    # Xử lý tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _render_paragraph(para, ctx)

    # Xử lý headers & footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for para in hf.paragraphs:
                    _render_paragraph(para, ctx)
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _render_paragraph(para, ctx)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_render_context(
    contract,          # Contract ORM object
    customer,          # Customer ORM object
    product=None,      # Product ORM object (optional)
    template=None,     # ContractTemplate ORM object (optional)
    db=None,           # SQLAlchemy Session (optional) — dùng để đọc Party B settings
) -> dict[str, str]:
    """
    Xây dựng render context từ các ORM object.
    Trả về dict {PLACEHOLDER_NAME: str_value} sẵn sàng cho render_docx().
    """
    ctx: dict[str, str] = {}

    # ── Bên A (Khách hàng) ──────────────────────────────────────────────────
    ctx["PARTY_A_NAME"]           = customer.legal_name or ""
    ctx["PARTY_A_ADDRESS"]        = customer.full_address() or ""
    # Tên đại diện: 2 biến thể
    #   PARTY_A_REPRESENTATIVE      = tên thuần (dùng ở chân ký, P168 đã có "Ông (Bà):" sẵn)
    #   PARTY_A_REPRESENTATIVE_FULL = "(Ông) Tên" / "(Bà) Tên" (dùng ở bảng thông tin trang đầu)
    _rep_name   = customer.representative_name or ""
    _rep_gender = getattr(customer, "representative_gender", None) or ""
    _prefix     = {"Nam": "(Ông)", "Nữ": "(Bà)"}.get(_rep_gender, "")
    ctx["PARTY_A_REPRESENTATIVE"]      = _rep_name
    ctx["PARTY_A_REPRESENTATIVE_FULL"] = f"{_prefix} {_rep_name}".strip() if _prefix else _rep_name
    ctx["PARTY_A_TITLE"]          = customer.representative_title or ""
    # Tên Bên A kèm xã/phường — vài mẫu cần dạng này
    ward_suffix = f", {customer.ward}" if customer.ward else ""
    ctx["PARTY_A_NAME_WARD"]      = f"{customer.legal_name or ''}{ward_suffix}"
    # Optional Party A fields (dùng để chèn dòng động vào bảng)
    ctx["PARTY_A_TAX_CODE"]       = customer.tax_code or ""
    ctx["PARTY_A_BANK_ACCOUNT"]   = customer.bank_account or ""
    ctx["PARTY_A_BANK_NAME"]      = customer.bank_name or ""
    ctx["PARTY_A_BENEFICIARY"]    = customer.beneficiary_name or ""
    ctx["PARTY_A_BUDGET_CODE"]    = getattr(customer, "budget_relation_code", None) or ""


    # ── Bên B (Viettel — cố định, nhưng đại diện hợp đồng cụ thể) ──────────
    # Đọc từ AppSetting nếu có db session, fallback về PARTY_B_DEFAULTS
    from app.services.settings_service import get_party_b_settings, PARTY_B_DEFAULTS
    if db is not None:
        _pbs = get_party_b_settings(db)
    else:
        _pbs = {
            "name":         PARTY_B_DEFAULTS["party_b_name"],
            "address":      PARTY_B_DEFAULTS["party_b_address"],
            "bank_account": PARTY_B_DEFAULTS["party_b_bank_account"],
            "bank_name":    PARTY_B_DEFAULTS["party_b_bank_name"],
            "beneficiary":  PARTY_B_DEFAULTS["party_b_beneficiary"],
            "tax_code":     PARTY_B_DEFAULTS["party_b_tax_code"],
            "title":        PARTY_B_DEFAULTS["party_b_title"],
            "representative":       PARTY_B_DEFAULTS["party_b_representative"],
            "authorization_number": PARTY_B_DEFAULTS["party_b_authorization_number"],
            "authorization_date":   PARTY_B_DEFAULTS["party_b_authorization_date"],
        }

    # Thông tin tổ chức Bên B (từ AppSetting, trước đây hardcode trong template Word)
    ctx["PARTY_B_NAME"]         = _pbs["name"]
    ctx["PARTY_B_ADDRESS"]      = _pbs["address"]
    ctx["PARTY_B_BANK_ACCOUNT"] = _pbs["bank_account"]
    ctx["PARTY_B_BANK_NAME"]    = _pbs["bank_name"]
    ctx["PARTY_B_BENEFICIARY"]  = _pbs["beneficiary"]
    ctx["PARTY_B_TAX_CODE"]     = _pbs["tax_code"]
    ctx["PARTY_B_TITLE"]        = _pbs["title"]

    # Đại diện ký & ủy quyền (ưu tiên từ contract object, fallback AppSetting)
    ctx["PARTY_B_REPRESENTATIVE"]       = contract.party_b_representative or _pbs["representative"]
    ctx["PARTY_B_AUTHORIZATION_NUMBER"] = contract.party_b_authorization_number or ""
    ctx["PARTY_B_AUTHORIZATION_DATE"]   = contract.party_b_authorization_date or ""

    # ── Số hợp đồng ─────────────────────────────────────────────────────────
    ctx["CONTRACT_NUMBER"] = contract.contract_number or ""
    # Địa điểm ký: lấy từ form, fallback về tên khách hàng
    ctx["SIGN_LOCATION"] = (contract.sign_location or "").strip() or (customer.legal_name or "")

    # ── Ngày ký ─────────────────────────────────────────────────────────────
    ctx["SIGN_DATE"] = format_date_vn(contract.sign_date)
    if contract.sign_date:
        parts = split_date_parts(contract.sign_date)
        ctx["SIGN_DATE_DAY"]   = parts["day"]
        ctx["SIGN_DATE_MONTH"] = parts["month"]
        ctx["SIGN_DATE_YEAR"]  = parts["year"]
    else:
        ctx["SIGN_DATE_DAY"] = ctx["SIGN_DATE_MONTH"] = ctx["SIGN_DATE_YEAR"] = ""

    # ── Ngày hiệu lực / kết thúc ────────────────────────────────────────────
    ctx["START_DATE"]  = format_date_vn(contract.start_date)
    ctx["END_DATE"]    = format_date_vn(contract.end_date)
    ctx["MONTH_COUNT"] = str(contract.month_count or "")

    # ── Số lượng / giá trị ──────────────────────────────────────────────────
    ctx["UNIT_COUNT"] = str(int(contract.unit_count or 1))

    qty = contract.quantity
    ctx["QUANTITY"] = _fmt_number(qty)

    ctx["UNIT_PRICE_PRE_VAT"] = _fmt_currency(contract.unit_price_pre_vat)
    ctx["VAT_RATE"]           = _fmt_vat(contract.vat_rate)
    ctx["UNIT_PRICE_VAT"]     = _fmt_currency(contract.unit_price_vat)
    ctx["TOTAL_AMOUNT"]       = _fmt_currency(contract.total_amount)

    # Số tiền bằng chữ — tự sinh hoặc dùng override
    if contract.amount_text:
        ctx["AMOUNT_TEXT"] = contract.amount_text
    elif contract.total_amount:
        try:
            ctx["AMOUNT_TEXT"] = amount_to_words(contract.total_amount)
        except Exception:
            ctx["AMOUNT_TEXT"] = ""
    else:
        ctx["AMOUNT_TEXT"] = ""

    # AMOUNT_TEXT_CHAN (biên bản thanh lý — thêm "chẵn")
    if contract.total_amount:
        try:
            ctx["AMOUNT_TEXT_CHAN"] = amount_to_words_chan(contract.total_amount)
        except Exception:
            ctx["AMOUNT_TEXT_CHAN"] = ctx["AMOUNT_TEXT"]
    else:
        ctx["AMOUNT_TEXT_CHAN"] = ""

    ctx["PRICING_DESCRIPTION"] = contract.pricing_description or ""

    # ── Nghiệm thu ───────────────────────────────────────────────────────────
    ctx["ACCEPTANCE_DATE"] = format_date_vn(contract.acceptance_date)
    if contract.acceptance_date:
        p = split_date_parts(contract.acceptance_date)
        ctx["ACCEPTANCE_DATE_DAY"]   = p["day"]
        ctx["ACCEPTANCE_DATE_MONTH"] = p["month"]
        ctx["ACCEPTANCE_DATE_YEAR"]  = p["year"]
    else:
        ctx["ACCEPTANCE_DATE_DAY"] = ctx["ACCEPTANCE_DATE_MONTH"] = ctx["ACCEPTANCE_DATE_YEAR"] = ""

    # ── Thanh lý ─────────────────────────────────────────────────────────────
    if contract.liquidation_date:
        p = split_date_parts(contract.liquidation_date)
        ctx["LIQUIDATION_DATE_DAY"]   = p["day"]
        ctx["LIQUIDATION_DATE_MONTH"] = p["month"]
        ctx["LIQUIDATION_DATE_YEAR"]  = p["year"]
        ctx["LIQUIDATION_DATE"]       = format_date_vn(contract.liquidation_date)
    else:
        ctx["LIQUIDATION_DATE_DAY"] = ctx["LIQUIDATION_DATE_MONTH"] = ctx["LIQUIDATION_DATE_YEAR"] = ""
        ctx["LIQUIDATION_DATE"] = ""

    return apply_template_mapping(ctx, template)


# ─── Internal helpers ────────────────────────────────────────────────────────

def _fmt_currency(value) -> str:
    """Format số tiền VNĐ: 24000000 → "24.000.000" (dấu chấm nghìn)."""
    if value is None:
        return ""
    try:
        n = int(Decimal(str(value)))
        return f"{n:,}".replace(",", ".")
    except Exception:
        return str(value)


def _fmt_number(value) -> str:
    """Format số đơn thuần không có đơn vị tiền tệ."""
    if value is None:
        return ""
    try:
        n = Decimal(str(value))
        if n == int(n):
            return str(int(n))
        return str(n)
    except Exception:
        return str(value)


def _fmt_vat(value) -> str:
    """Format VAT: Decimal('10.00') → '10%', Decimal('0') → '0%'."""
    if value is None:
        return ""
    try:
        n = Decimal(str(value))
        if n == int(n):
            return f"{int(n)}%"
        return f"{n}%"
    except Exception:
        return str(value)


def _replace_in_text(text: str, ctx: dict) -> str:
    """Thay thế tất cả {{KEY}} trong text bằng ctx[KEY], giữ nguyên nếu KEY không có."""
    def replacer(m):
        return ctx.get(m.group(1), m.group(0))   # giữ {{KEY}} nếu không có mapping
    return _PH_RE.sub(replacer, text)


def _render_paragraph(para, ctx: dict) -> None:
    """
    Render tất cả placeholder trong một paragraph.
    Dùng chiến lược 2 lớp để xử lý cả split-run.
    """
    if not para.runs:
        return

    full_text = para.text
    if "{{" not in full_text:
        return

    # ── Lớp 1: Single-run replace ──────────────────────────────────────────
    for run in para.runs:
        if "{{" in run.text:
            run.text = _replace_in_text(run.text, ctx)

    # ── Lớp 2: Merge fallback nếu còn split-run ────────────────────────────
    if "{{" in para.text:
        _merge_and_replace(para, ctx)


def _merge_and_replace(para, ctx: dict) -> None:
    """
    Fallback cho split-run: merge toàn bộ runs vào run đầu tiên,
    replace text, xóa runs thừa.
    Formatting của paragraph (bold/italic của run đầu) được giữ lại.
    Đây là trade-off chấp nhận được cho V1 — contract templates thường
    dùng cùng 1 format cho toàn bộ đoạn có placeholder.
    """
    runs = para.runs
    if not runs:
        return

    # Build full text
    full_text = "".join(r.text for r in runs)
    rendered  = _replace_in_text(full_text, ctx)

    # Set text vào run đầu tiên
    runs[0].text = rendered

    # Xóa các run còn lại (từ index 1 trở đi) khỏi XML
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
