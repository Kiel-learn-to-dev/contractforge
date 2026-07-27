#!/usr/bin/env python3
"""Generate the neutral sample DOCX templates shipped with the public repo.

The repository deliberately ignores ``ContractForge/assets/default_templates/*.docx``
because a real installation's templates carry that organization's letterhead and
wording. But one of them is not user content: the acceptance form (08A) is
rendered on demand by ``form08a_service`` and the feature simply breaks on a
fresh clone if no file is present.

This script builds a fictional, structurally-equivalent stand-in and writes it as
``sample_form_08a.docx``, which .gitignore explicitly un-ignores. Run it whenever
the placeholder set changes:

    python scripts/make_sample_templates.py

It is a generator, not a runtime dependency — the app only reads the .docx.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = (Path(__file__).resolve().parent.parent
       / "ContractForge" / "assets" / "default_templates" / "sample_form_08a.docx")


def _para(doc, text="", *, bold=False, italic=False, align=None, size=11):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def build() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Header block ────────────────────────────────────────────────────────
    head = doc.add_table(rows=1, cols=3)
    head.alignment = WD_TABLE_ALIGNMENT.CENTER
    head.rows[0].cells[0].text = "{{BUDGET_UNIT_NAME}}"
    head.rows[0].cells[1].text = ""
    head.rows[0].cells[2].text = "Mẫu số 08a"

    _para(doc)
    _para(doc, "BẢNG XÁC ĐỊNH GIÁ TRỊ KHỐI LƯỢNG CÔNG VIỆC HOÀN THÀNH",
          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    _para(doc, "THEO HỢP ĐỒNG", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    _para(doc)

    _para(doc, "Đơn vị sử dụng ngân sách: {{BUDGET_UNIT_NAME}}")
    _para(doc, "Mã đơn vị quan hệ ngân sách: {{BUDGET_UNIT_CODE}}")
    _para(doc, "Căn cứ hợp đồng số: {{CONTRACT_NUMBER}} ngày {{SIGN_DATE}}")
    _para(doc, "Bên A: {{PARTY_A_NAME}}")
    _para(doc)

    # ── Work table ──────────────────────────────────────────────────────────
    table = doc.add_table(rows=4, cols=6)
    table.style = "Table Grid"
    headers = ["STT", "Nội dung công việc", "Đơn vị tính",
               "Số lượng", "Đơn giá", "Thành tiền"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    row = table.rows[1].cells
    row[0].text = "1"
    row[1].text = "{{WORK_CONTENT}}"
    row[2].text = "Tháng"
    row[3].text = "{{QUANTITY}}"
    row[4].text = "{{UNIT_PRICE_VAT}}"
    row[5].text = "{{TOTAL_AMOUNT}}"

    total = table.rows[2].cells
    total[0].text = ""
    total[1].text = "Tổng cộng"
    total[5].text = "{{TOTAL_AMOUNT}}"

    note = table.rows[3].cells
    note[1].text = "(Giá trên đã bao gồm thuế giá trị gia tăng)"

    _para(doc)
    _para(doc,
          "Ngày {{ACCEPTANCE_DATE_DAY}} tháng {{ACCEPTANCE_DATE_MONTH}} "
          "năm {{ACCEPTANCE_DATE_YEAR}}",
          italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc)

    # ── Signature block ─────────────────────────────────────────────────────
    sign = doc.add_table(rows=2, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign.rows[0].cells[0].text = "{{PARTY_A_TITLE_UPPER}}"
    sign.rows[0].cells[1].text = "{{PARTY_B_TITLE_UPPER}}"
    sign.rows[1].cells[0].text = "{{PARTY_A_REPRESENTATIVE}}"
    sign.rows[1].cells[1].text = "{{PARTY_B_REPRESENTATIVE}}"
    for row_ in sign.rows:
        for cell in row_.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"Đã ghi {OUT} ({OUT.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
