"""Integration coverage for product/template-driven quotations (Task 7)."""

from __future__ import annotations

from datetime import date
from io import BytesIO

import pytest
from docx import Document
from starlette.requests import Request


def test_quotation_choices_include_only_active_products_and_templates(
    db_session, make_product, make_template,
):
    from app.routers.quotation import (
        list_quotation_products,
        list_quotation_templates,
    )

    available_template = make_template(name="Mẫu báo giá chung", contract_type="Báo giá")
    available_product = make_product(code="QUOTE-AVAILABLE", name="Gói có thể báo giá")
    make_product(code="QUOTE-HIDDEN", name="Gói đã tắt", is_active=False)
    make_template(name="Mẫu đã tắt", is_active=False, contract_type="Báo giá")

    products = list_quotation_products(db_session)
    templates = list_quotation_templates(db_session)

    assert [product.id for product in products] == [available_product.id]
    assert [template.id for template in templates] == [available_template.id]


def test_product_price_and_vat_drive_quotation_amounts(db_session, make_product):
    from app.routers.quotation import product_price_with_vat

    product = make_product(default_price=250_000, default_vat_rate=8)

    assert product_price_with_vat(product) == 270_000

    legacy_rounding_case = make_product(
        code="ROUNDING-CASE",
        default_price=440_909,
        default_vat_rate=8,
    )
    assert product_price_with_vat(legacy_rounding_case) == 476_182


def test_resolve_selection_rejects_inactive_records(
    db_session, make_product, make_template,
):
    from app.routers.quotation import resolve_quotation_selection

    product = make_product(is_active=False)
    template = make_template(contract_type="Báo giá")

    with pytest.raises(ValueError, match="Sản phẩm"):
        resolve_quotation_selection(db_session, product.id, template.id)

    unpriced_product = make_product(code="NO-PRICE", default_price=None)
    with pytest.raises(ValueError, match="đơn giá"):
        resolve_quotation_selection(db_session, unpriced_product.id, template.id)


def test_context_uses_selected_product_configuration(db_session, make_customer, make_product):
    from app.routers.quotation import _build_context

    customer = make_customer(legal_name="Khách hàng Mẫu")
    product = make_product(
        code="SERVICE-CUSTOM",
        name="Dịch vụ Tùy chỉnh",
        default_price=250_000,
        default_vat_rate=8,
    )

    context = _build_context(customer, product, date(2026, 7, 11), 12)

    assert context["PRODUCT_CODE"] == "SERVICE-CUSTOM"
    assert context["PRODUCT_NAME"] == "Dịch vụ Tùy chỉnh"
    assert context["UNIT_PRICE_PRE_VAT"] == "250.000"
    assert context["VAT_RATE"] == "8%"
    assert context["UNIT_PRICE_VAT"] == "270.000"
    assert context["TOTAL_AMOUNT"] == "3.240.000"


def test_generate_uses_user_selected_product_and_template(
    db_session, make_customer, make_product, make_template, tmp_path, monkeypatch,
):
    import app.routers.quotation as quotation

    customer = make_customer(short_name="KH Mẫu")
    product = make_product(
        code="SERVICE-CUSTOM",
        default_price=250_000,
        default_vat_rate=8,
    )
    template_path = tmp_path / "quotation.docx"
    template_path.write_bytes(b"placeholder")
    template = make_template(file_path=str(template_path), file_name="quotation.docx",
                             contract_type="Báo giá")
    captured = {}

    def fake_render(path, context):
        captured["path"] = path
        captured["context"] = context
        return b"rendered-docx"

    monkeypatch.setattr(quotation, "SessionLocal", lambda: db_session)
    monkeypatch.setattr(quotation, "render_docx", fake_render)
    request = Request({
        "type": "http",
        "method": "POST",
        "path": "/quotation/generate",
        "query_string": b"",
        "headers": [],
    })

    response = quotation.quotation_generate(
        request=request,
        product_id=product.id,
        template_id=template.id,
        customer_id=customer.id,
        quote_date="2026-07-11",
        month_count=12,
    )

    assert response.body == b"rendered-docx"
    assert captured["path"] == str(template_path)
    assert captured["context"]["PRODUCT_CODE"] == "SERVICE-CUSTOM"


def test_form_renders_product_and_template_choices(
    db_session, make_product, make_template, monkeypatch,
):
    import app.routers.quotation as quotation

    make_product(code="SERVICE-CUSTOM", name="Dịch vụ Tùy chỉnh")
    make_template(code="TPL-QUOTE", name="Mẫu báo giá tùy chỉnh", contract_type="Báo giá")
    monkeypatch.setattr(quotation, "SessionLocal", lambda: db_session)
    request = Request({
        "type": "http",
        "method": "GET",
        "path": "/quotation",
        "query_string": b"",
        "headers": [],
    })

    response = quotation.quotation_form(request)
    body = response.body.decode("utf-8")

    assert "Dịch vụ Tùy chỉnh" in body
    assert "Mẫu báo giá tùy chỉnh" in body
    assert 'name="product_id"' in body
    assert 'name="template_id"' in body
    assert 'data-duration="12"' in body


def test_custom_placeholder_mapping_is_applied(db_session, make_template):
    from app.services.docx_renderer import apply_template_mapping

    template = make_template()
    template.field_mapping = {
        "TEN_SAN_PHAM": "product.name",
        "NGAY_BAO_GIA": "quotation.date",
    }
    db_session.commit()

    mapped = apply_template_mapping(
        {
            "PRODUCT_NAME": "Dịch vụ Tùy chỉnh",
            "QUOTE_DATE": "11/07/2026",
        },
        template,
    )

    assert mapped["TEN_SAN_PHAM"] == "Dịch vụ Tùy chỉnh"
    assert mapped["NGAY_BAO_GIA"] == "11/07/2026"


def test_custom_placeholder_mapping_renders_into_real_docx(db_session, make_template, tmp_path):
    from app.services.docx_renderer import apply_template_mapping, render_docx

    source = tmp_path / "custom-quotation.docx"
    document = Document()
    document.add_paragraph("Sản phẩm: {{TEN_SAN_PHAM}}")
    document.save(source)
    template = make_template(file_path=str(source), file_name=source.name)
    template.field_mapping = {"TEN_SAN_PHAM": "product.name"}
    db_session.commit()

    context = apply_template_mapping({"PRODUCT_NAME": "Dịch vụ Tùy chỉnh"}, template)
    rendered = Document(BytesIO(render_docx(str(source), context)))

    assert rendered.paragraphs[0].text == "Sản phẩm: Dịch vụ Tùy chỉnh"


# ─── Hồi quy: báo giá không được sinh ra file hợp đồng ────────────────────────

def test_contract_templates_never_appear_in_the_quotation_picker(
    db_session, make_template,
):
    """Ô chọn mẫu ở trang báo giá chỉ được liệt kê mẫu phân loại 'Báo giá'.

    Lỗi cũ: hàm này trả về mọi mẫu đang bật. Mẫu hợp đồng lọt vào danh sách, và
    đoạn JS tự chọn `product.default_template_id` — vốn trỏ vào mẫu HỢP ĐỒNG của
    sản phẩm — khiến chọn sản phẩm xong là ô mẫu tự nhảy sang mẫu hợp đồng.
    """
    from app.routers.quotation import list_quotation_templates

    quote_tpl = make_template(code="TPL-BG", name="Mẫu báo giá",
                              contract_type="Báo giá")
    make_template(code="TPL-HD", name="Mẫu chuẩn hợp đồng", contract_type="Hợp đồng")
    make_template(code="TPL-NT", name="Biểu nghiệm thu", contract_type="Nghiệm thu")

    assert [t.id for t in list_quotation_templates(db_session)] == [quote_tpl.id]


def test_generating_with_a_contract_template_is_refused(
    db_session, make_product, make_template,
):
    """Chặn ở tầng server, không chỉ ở dropdown.

    Form còn mở ở tab khác, nút Back, hoặc mẫu bị đổi nhãn sau khi trang đã tải
    đều gửi lên được template_id của mẫu hợp đồng. Nếu không chặn, người dùng
    nhận về một bản hợp đồng mang tên "BaoGia_....docx" và chỉ phát hiện sau khi
    đã gửi cho khách.
    """
    from app.routers.quotation import resolve_quotation_selection

    product = make_product(code="SP-BG", default_price=100_000)
    contract_tpl = make_template(code="TPL-HD2", name="Mẫu chuẩn hợp đồng",
                                 contract_type="Hợp đồng")

    with pytest.raises(ValueError, match="không phải mẫu Báo giá"):
        resolve_quotation_selection(db_session, product.id, contract_tpl.id)


def test_product_default_template_is_a_contract_template_not_a_quotation_one(
    db_session, make_product, make_template,
):
    """Ghim lại gốc rễ của lỗi: default_template_id của sản phẩm là mẫu HỢP ĐỒNG.

    Nếu sau này có ai định dùng lại trường đó để gợi ý mẫu cho trang báo giá thì
    test này nhắc rằng nó mang ý nghĩa khác.
    """
    from app.routers.quotation import list_quotation_templates

    contract_tpl = make_template(code="TPL-HD3", contract_type="Hợp đồng")
    product = make_product(code="SP-BG2", default_template_id=contract_tpl.id)

    assert product.default_template_id == contract_tpl.id
    assert contract_tpl.id not in {t.id for t in list_quotation_templates(db_session)}
